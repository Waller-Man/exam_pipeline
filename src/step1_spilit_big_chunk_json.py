# -*- coding: utf-8 -*-
import os
import re
import json
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from dotenv import load_dotenv
from docx import Document
from volcenginesdkarkruntime import Ark

load_dotenv()

BASE_URL = os.getenv("ARK_BASE_URL", "https://ark.cn-beijing.volces.com/api/v3")
API_KEY = os.getenv("ARK_API_KEY")
TEXT_MODEL = os.getenv("ARK_TEXT_MODEL")

# 你要求：分段尽量 < 5000 词（这里做近似估算：中文按“汉字数”，英文按“word数”）
MAX_WORDS = 5000

# 允许极少数“单个材料块”本身就超过5000，此时不切碎（会略超）
HARD_MAX_WORDS = 6500

# ----------- 识别规则（轻量） -----------
MAJOR_RE = re.compile(r"^[一二三四五六七八九十]+、")  # 一、二、...
QTYPE_RE = re.compile(r"(单项选择题|多项选择题|判断题|填空题|简答题|综合题|操作题|选择题|阅读理解|完形填空|语法填空|七选五|任务型阅读)$")
QSTART_RE = re.compile(r"^(?:小题\s*)?\(?\s*(\d{1,3})\s*\)?\s*[\.、\)]\s*")  # 1. / 1、 / (1) / 小题1

# 材料/文章起始提示（中英）
MATERIAL_HINT_RE = re.compile(
    r"(阅读理解|完形填空|阅读材料|短文|材料|Passage|Text\s*\d+|Cloze|Reading\s+Comprehension|Directions:|Read\s+the\s+following)",
    re.IGNORECASE
)

def estimate_words(text: str) -> int:
    """近似词数：中文按汉字数，英文按word数"""
    if not text:
        return 0
    cn = len(re.findall(r"[\u4e00-\u9fff]", text))
    en = len(re.findall(r"[A-Za-z0-9]+(?:'[A-Za-z0-9]+)?", text))
    # 取两者相加，作为“近似词/字数”
    return cn + en

def norm_line(t: str) -> str:
    t = (t or "").replace("\u3000", " ").strip()
    t = re.sub(r"[ \t]+", " ", t)
    return t.strip()

def is_heading(line: str) -> bool:
    s = (line or "").strip()
    return bool(s) and (MAJOR_RE.match(s) or QTYPE_RE.search(s))

def is_qstart(line: str) -> bool:
    return QSTART_RE.match((line or "").strip()) is not None

def looks_like_material_intro(line: str) -> bool:
    s = (line or "").strip()
    if not s:
        return False
    # 材料提示行通常不是“题号行”
    return bool(MATERIAL_HINT_RE.search(s)) and (not is_qstart(s))

def read_docx_lines(docx_path: Path) -> List[str]:
    doc = Document(str(docx_path))
    lines: List[str] = []
    for p in doc.paragraphs:
        t = norm_line(p.text)
        if t:
            lines.append(t)
    # 表格：逐行拼接
    for table in doc.tables:
        for row in table.rows:
            cells = [norm_line(c.text) for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                lines.append(" | ".join(cells))
    return lines

# ----------------- 核心：轻量“块”构建（不切碎材料+小题） -----------------

def build_blocks(lines: List[str]) -> List[str]:
    """
    产出 blocks，每个 block 是一段文本（不会把“材料 + 紧跟的小题串”拆开）：
    - 标题行单独作为 block
    - 普通小题：从题号行开始到下一个题号/标题之前
    - 材料题：材料引导/原文段 + 后续连续小题，合成一个 block（交给模型去切题为“大题”）
    """
    blocks: List[str] = []
    i = 0
    n = len(lines)

    while i < n:
        s = (lines[i] or "").strip()
        if not s:
            i += 1
            continue

        # 标题单独块
        if is_heading(s):
            blocks.append(s)
            i += 1
            continue

        # 材料引导：把“材料段 + 后续小题串”合成一个块
        if looks_like_material_intro(s):
            buf = [s]
            i += 1

            # 吸收后续“非题号”的原文/说明行，直到遇到第一道小题或标题
            while i < n:
                t = (lines[i] or "").strip()
                if not t:
                    i += 1
                    continue
                if is_heading(t) or is_qstart(t) or looks_like_material_intro(t):
                    break
                buf.append(t)
                i += 1

            # 吸收连续小题（直到遇到标题或下一个材料引导）
            # 注意：这里不深度解析每小题边界，只保证它们跟材料在同块里即可
            while i < n:
                t = (lines[i] or "").strip()
                if not t:
                    i += 1
                    continue
                if is_heading(t) or looks_like_material_intro(t):
                    break
                buf.append(t)
                i += 1

            blocks.append("\n".join(buf).strip())
            continue

        # 普通：从当前位置开始，直到下一个标题/材料引导（保守）
        buf = [s]
        i += 1
        while i < n:
            t = (lines[i] or "").strip()
            if not t:
                i += 1
                continue
            if is_heading(t) or looks_like_material_intro(t):
                break
            # 如果当前是小题起始行，且buf里也已经有一个小题起始，说明可能进入下一题，保守切块
            if is_qstart(t) and any(is_qstart(x) for x in buf):
                break
            buf.append(t)
            i += 1
        blocks.append("\n".join(buf).strip())

    return blocks

def pack_blocks_to_chunks(blocks: List[str], max_words: int) -> List[str]:
    """
    把 blocks 打包成 chunk，每个 chunk 尽量 < max_words；
    但单个 block 若本身超限，不切碎（最多允许到 HARD_MAX_WORDS）
    """
    chunks: List[str] = []
    cur: List[str] = []
    cur_words = 0

    def flush():
        nonlocal cur, cur_words
        if cur:
            chunks.append("\n\n".join(cur).strip())
        cur = []
        cur_words = 0

    for b in blocks:
        w = estimate_words(b)
        if not b:
            continue

        # 单块特别大：独占一段（不切碎）
        if w >= max_words:
            flush()
            chunks.append(b)
            continue

        if cur and (cur_words + w) > max_words:
            flush()

        cur.append(b)
        cur_words += w

    flush()
    return chunks

# ----------------- Prompt：强制模型按“材料+小题=一道大题”出JSON -----------------

def build_prompt(
    chapter: str,
    part_no: int,
    part_total: int,
    prev_section: Optional[Dict[str, Any]],
    text: str
) -> str:
    prev_section = prev_section or {"section_index": None, "section_title": None}

    schema = {
        "sections": [
            {
                "section_index": 1,
                "section_title": "字符串",
                "section_inferred": False,
                "questions": [
                    {
                        "question_number": 1,  # 或 "11-15"
                        "raw_text": "string",
                        "has_figure": False,
                        "figure_note": None,
                        "errors": [],
                        "confidence": 0.0
                    }
                ]
            }
        ],
        "errors": [],
        "confidence": 0.0
    }

    return f"""
你将收到章节《{chapter}》的一段文本（第 {part_no}/{part_total} 段）。请抽取题目并结构化为 JSON。
必须严格只输出 JSON（不要解释、不要 markdown、不要任何多余字符）。

【最关键要求：材料题合并成一道大题】
- 若文本包含“阅读理解/完形填空/材料/Passage/Text/Cloze”等原文，并且后面紧跟若干选择题小题：
  必须把“原文 + 紧跟的小题（含选项）”合并为 questions 中的【一个】对象（当作一道大题）。
- 合并后的 question_number 用字符串范围，例如 "11-15"（表示该材料对应11~15小题）。
- 合并后的 raw_text 必须完整包含：
  【原文】
  ...原文...
  【小题11】...A...B...C...D...
  【小题12】...
  （不要拆成多条 questions）

通用要求：
1) 识别大题/题型标题（如“一、...”“单项选择题”“阅读理解”等）组织为 sections。
2) 普通选择题：raw_text 中选项尽量一行一个：A. ...\\nB. ...\\nC. ...\\nD. ...
3) 含图：若题干出现“如图/见图/图x-xx/[含图]”，has_figure=true，figure_note 简短，不输出图片内容。
4) 如果本段开头没有新标题，沿用 prev_section，并设置 section_inferred=true。
5) 如果发现明显残缺（末尾被截断），在该题 errors 中加入 "TRUNCATED"。

输出 JSON 结构必须与下方示例字段一致（可多题、多 section）：
{json.dumps(schema, ensure_ascii=False)}

prev_section:
{json.dumps(prev_section, ensure_ascii=False)}

文本如下：
{text}
""".strip()

def safe_json_loads(s: str) -> dict:
    s = (s or "").strip()
    try:
        return json.loads(s)
    except Exception:
        l = s.find("{")
        r = s.rfind("}")
        if l != -1 and r != -1 and r > l:
            return json.loads(s[l:r+1])
        raise

def ark_chat_json(client: Ark, prompt: str) -> dict:
    # 尽量强制json_object，不支持则降级
    try:
        resp = client.chat.completions.create(
            model=TEXT_MODEL,
            temperature=0.1,
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
        )
        return safe_json_loads(resp.choices[0].message.content or "")
    except Exception:
        resp = client.chat.completions.create(
            model=TEXT_MODEL,
            temperature=0.1,
            messages=[{"role": "user", "content": prompt}],
        )
        return safe_json_loads(resp.choices[0].message.content or "")

def flatten_questions(chapter: str, pages: List[dict]) -> List[dict]:
    out: List[dict] = []
    seq = 0
    for page in pages:
        meta = page.get("_meta", {})
        page_no = meta.get("page_no")
        image = meta.get("image")

        for sec in page.get("sections", []):
            sec_idx = sec.get("section_index")
            sec_title = sec.get("section_title")
            sec_inferred = bool(sec.get("section_inferred", False))
            sec_key = f"IDX:{sec_idx}|{sec_title}"

            for q in sec.get("questions", []):
                seq += 1
                out.append({
                    "uid": f"Q{seq:05d}",
                    "chapter": chapter,
                    "page_no": page_no,
                    "image": image,
                    "section_index": sec_idx,
                    "section_title": sec_title,
                    "section_key": sec_key,
                    "question_number": q.get("question_number"),
                    "raw_text": (q.get("raw_text") or "").strip(),
                    "has_figure": bool(q.get("has_figure", False)),
                    "figure_note": q.get("figure_note", None),
                    "errors": q.get("errors", []),
                    "confidence": float(q.get("confidence", 0.85)),
                    "section_inferred": sec_inferred,
                })
    return out

def main(chapter: str, docx_path: str):
    if not API_KEY:
        raise RuntimeError("缺少 ARK_API_KEY（请检查 .env）")
    if not TEXT_MODEL:
        raise RuntimeError("缺少 ARK_TEXT_MODEL（请检查 .env）")

    in_path = Path(docx_path)
    if not in_path.exists():
        raise FileNotFoundError(f"找不到：{in_path.resolve()}")

    out_dir = Path("out") / chapter
    out_dir.mkdir(parents=True, exist_ok=True)
    debug_dir = out_dir / "_debug_step1_docx_light"
    debug_dir.mkdir(parents=True, exist_ok=True)

    print(f"[阶段] 1/5 读取 docx：{in_path}", flush=True)
    lines = read_docx_lines(in_path)

    print("[阶段] 2/5 轻量分块：保证材料+小题不被切碎...", flush=True)
    blocks = build_blocks(lines)

    print("[阶段] 3/5 打包分段：尽量 < 5000 词/字...", flush=True)
    chunks = pack_blocks_to_chunks(blocks, MAX_WORDS)
    print(f"[信息] 分段数：{len(chunks)}", flush=True)

    client = Ark(base_url=BASE_URL, api_key=API_KEY)

    pages: List[dict] = []
    prev_section: Optional[Dict[str, Any]] = None

    for idx, chunk in enumerate(chunks, start=1):
        w = estimate_words(chunk)
        if w > HARD_MAX_WORDS:
            # 仍不切碎，只记录警告
            (debug_dir / f"part_{idx:03d}_too_long.txt").write_text(
                f"estimated_words={w}\n\n{chunk}", encoding="utf-8"
            )

        print(f"[阶段] 4/5 调用模型抽题：{idx}/{len(chunks)}（估算词/字={w}）...", flush=True)
        prompt = build_prompt(chapter, idx, len(chunks), prev_section, chunk)

        obj = ark_chat_json(client, prompt)
        (debug_dir / f"part_{idx:03d}_raw.json").write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")

        # 更新 prev_section：取本段最后一个 section
        secs = obj.get("sections", [])
        if isinstance(secs, list) and secs:
            last = secs[-1]
            prev_section = {
                "section_index": last.get("section_index"),
                "section_title": last.get("section_title"),
            }

        pages.append({
            "_meta": {"page_no": idx, "image": f"{in_path.name}#part{idx}"},
            "sections": obj.get("sections", []),
            "errors": obj.get("errors", []),
            "confidence": obj.get("confidence", 0.85),
        })

    questions = flatten_questions(chapter, pages)

    chapter_json = {
        "chapter_name": chapter,
        "source_type": "docx",
        "source_file": in_path.name,
        "model": TEXT_MODEL,
        "pages": pages,        # step2 用于“原文整理版”
        "questions": questions, # step2 用于“抽样改编”
        "errors": [],
    }

    print("[阶段] 5/5 写出章节总 JSON...", flush=True)
    out_path = out_dir / f"{chapter}.json"
    out_path.write_text(json.dumps(chapter_json, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[完成] 输出：{out_path.resolve()}", flush=True)
    print(f"[调试] debug：{debug_dir.resolve()}", flush=True)

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("用法：python src\\step1_docx_build_chapter_json.py ch01 path\\to\\xxx.docx")
        raise SystemExit(1)
    main(sys.argv[1], sys.argv[2])
