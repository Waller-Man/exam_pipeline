# -*- coding: utf-8 -*-
import os
import re
import json
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from dotenv import load_dotenv
from docx import Document
from volcenginesdkarkruntime import Ark

load_dotenv()

BASE_URL = os.getenv("ARK_BASE_URL", "https://ark.cn-beijing.volces.com/api/v3")
API_KEY = os.getenv("ARK_API_KEY")
TEXT_MODEL = os.getenv("ARK_TEXT_MODEL") or os.getenv("ARK_VISION_MODEL")

# docx很长时的自动分段阈值（固定内置，不需要你调参）
_MAX_CHUNK_CHARS = 5000

# 如果一段失败，自动二分的最小长度阈值（过短就不再拆）
_MIN_SPLIT_CHARS = 2500

# 单段最多拆分次数，防止极端情况下无限拆
_MAX_SPLIT_PER_ORIGINAL_CHUNK = 10


# ---------------- JSON 解析兜底 ----------------

def _safe_json_loads(s: str) -> dict:
    """
    尝试解析 JSON。
    允许 JSON 前后夹杂少量文字：截取最外层 {} 再解析。
    """
    s = (s or "").strip()
    try:
        return json.loads(s)
    except Exception:
        l = s.find("{")
        r = s.rfind("}")
        if l != -1 and r != -1 and r > l:
            return json.loads(s[l:r + 1])
        raise


def _save_debug_text(out_dir: Path, name: str, text: str) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / name).write_text(text or "", encoding="utf-8")


def _fix_json_via_model(client: Ark, bad_text: str) -> dict:
    """
    让模型把“接近JSON但不合法”的文本修复成合法 JSON 并返回 dict。
    """
    prompt = f"""
下面内容应该是一个 JSON，但不是合法 JSON（常见原因：字符串内有未转义换行/引号，或格式缺逗号/括号）。
请你将其修复为“严格合法 JSON”，并且只输出 JSON 本体，不要解释，不要 markdown。

关键要求：
- JSON 字符串内不要出现真实换行，必须用 \\n
- JSON 字符串内若包含引号 " 必须转义为 \\\\"
- 确保所有括号、逗号、引号闭合

待修复内容：
{bad_text}
""".strip()

    resp = client.chat.completions.create(
        model=TEXT_MODEL,
        temperature=0.1,
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"},
    )

    return _safe_json_loads(resp.choices[0].message.content)


# ---------------- docx 读取与分段 ----------------

def _read_docx_as_lines(docx_path: Path) -> List[str]:
    """读取 docx：段落 + 表格，转换为行文本"""
    doc = Document(str(docx_path))
    lines: List[str] = []

    def norm(t: str) -> str:
        t = (t or "").replace("\u3000", " ").strip()
        t = re.sub(r"[ \t]+", " ", t)
        return t.strip()

    # 段落
    for p in doc.paragraphs:
        t = norm(p.text)
        if t:
            lines.append(t)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [norm(c.text) for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                lines.append(" | ".join(cells))

    return lines


_MAJOR_RE = re.compile(r"^[一二三四五六七八九十]+、")
_QSTART_RE = re.compile(r"^\d+\s*[\.、]")
_QTYPE_HINT_RE = re.compile(r"(单项选择题|多项选择题|判断题|填空题|简答题|综合题|操作题|选择题)$")


def _is_boundary(line: str) -> bool:
    """裁切安全边界：大题/题型/小题起始行"""
    s = (line or "").strip()
    if not s:
        return False
    if _MAJOR_RE.match(s):
        return True
    if _QSTART_RE.match(s):
        return True
    if _QTYPE_HINT_RE.search(s):
        return True
    return False


def _chunk_lines(lines: List[str], max_chars: int) -> List[str]:
    """
    按 max_chars 做裁切；尽量只在“边界行”前切断，避免把一道题切成两段。
    返回：chunk_text 列表
    """
    chunks: List[str] = []
    cur: List[str] = []
    cur_len = 0
    last_safe_cut = 0  # cur 内“安全切点”的行下标

    for ln in lines:
        if _is_boundary(ln):
            last_safe_cut = len(cur)

        cur.append(ln)
        cur_len += len(ln) + 1

        if cur_len >= max_chars and len(cur) > 10:
            cut = last_safe_cut if last_safe_cut > 0 else len(cur)
            chunk = "\n".join(cur[:cut]).strip()
            if chunk:
                chunks.append(chunk)

            rest = cur[cut:]
            cur = rest
            cur_len = sum(len(x) + 1 for x in cur)

            # 重算 safe cut
            last_safe_cut = 0
            for i, x in enumerate(cur):
                if _is_boundary(x):
                    last_safe_cut = i

    tail = "\n".join(cur).strip()
    if tail:
        chunks.append(tail)
    return chunks


# ---------------- Prompt 与规范化 ----------------

def _build_extract_prompt(
    chapter_name: str,
    part_no: int,
    part_total: int,
    prev_section_index: Optional[int],
    prev_section_title: Optional[str],
    text: str,
) -> str:
    prev_hint = {
        "section_index": prev_section_index,
        "section_title": prev_section_title,
    }

    return f"""
你将收到章节《{chapter_name}》的一部分题库文本（第 {part_no}/{part_total} 段）。请从中抽取题目并结构化。
必须只输出 JSON（不要输出解释、不要 markdown、不要任何多余字符）。

非常重要（避免JSON解析失败）：
- JSON 字符串里禁止出现真实换行；必须用 \\n 表示换行。
- JSON 字符串里如果出现引号 "，必须转义为 \\\\"。

抽取规则：
1) 识别“分组/大题/题型”标题（如“一、……”“单项选择题”“操作题(……)”等），组织为 sections。
2) 识别小题起始（如 1. / 1、 / 128. 等），题干与选项合并为 raw_text。
3) raw_text 用 \\n 分行：题干一段；选项尽量一行一个（A. ...\\nB. ...）。
4) 若题干提到图（如“如图”“图1-2-13”“见图”），has_figure=true，figure_note 简短；不要输出图片内容。
5) 如果本段开头没有出现新的分组/题型标题，则默认沿用上一段 prev_section（见下方），并把 section_inferred=true。
6) 不要编造缺失内容；若发现明显残缺（如本段末尾截断），在该题 errors 中加入 "TRUNCATED"。

输出 JSON 结构（必须完全一致）：
{{
  "sections": [
    {{
      "section_index": 1,
      "section_title": "字符串",
      "section_inferred": false,
      "questions": [
        {{
          "question_number": 1,
          "raw_text": "题干...\\nA....\\nB....\\nC....\\nD....",
          "has_figure": false,
          "figure_note": null,
          "errors": [],
          "confidence": 0.0
        }}
      ]
    }}
  ],
  "errors": [],
  "confidence": 0.0
}}

prev_section（如本段无新标题可沿用）：
{json.dumps(prev_hint, ensure_ascii=False)}

本段原始文本：
{text}
""".strip()


def _normalize_section(sec: dict, fallback_idx: Optional[int], fallback_title: Optional[str]) -> dict:
    """补齐缺失的 section_index/title，并标准化 question 字段"""
    sec_idx = sec.get("section_index", None)
    sec_title = sec.get("section_title", None)

    inferred = bool(sec.get("section_inferred", False))

    if sec_idx is None and fallback_idx is not None:
        sec["section_index"] = fallback_idx
        inferred = True

    if (sec_title is None or str(sec_title).strip() == "") and fallback_title:
        sec["section_title"] = fallback_title
        inferred = True

    sec["section_inferred"] = inferred

    qs = sec.get("questions", [])
    if not isinstance(qs, list):
        qs = []

    new_qs = []
    for q in qs:
        if not isinstance(q, dict):
            continue
        q.setdefault("question_number", None)
        q.setdefault("raw_text", "")
        q.setdefault("has_figure", False)
        q.setdefault("figure_note", None)
        q.setdefault("errors", [])
        q.setdefault("confidence", 0.85)
        new_qs.append(q)
    sec["questions"] = new_qs

    return sec


# ---------------- 主流程 ----------------

def main(chapter_name: str, docx_path: str):
    if not API_KEY or not TEXT_MODEL:
        raise RuntimeError("缺少环境变量：ARK_API_KEY 或 ARK_TEXT_MODEL（请检查 .env）")

    chapter_name = chapter_name.strip()
    in_path = Path(str(docx_path).strip())
    if not in_path.exists() or in_path.suffix.lower() != ".docx":
        raise FileNotFoundError(f"输入文件不存在或不是 .docx：{in_path.resolve()}")

    out_dir = Path("out") / chapter_name
    out_dir.mkdir(parents=True, exist_ok=True)
    debug_dir = out_dir / "_debug"
    out_json_path = out_dir / f"{chapter_name}.json"

    print(f"[阶段] 1/5 读取 docx：{in_path}", flush=True)
    lines = _read_docx_as_lines(in_path)

    print("[阶段] 2/5 自动裁切为多段（避免上下文过长）...", flush=True)
    chunks = _chunk_lines(lines, max_chars=_MAX_CHUNK_CHARS)
    print(f"[信息] 初始分段数：{len(chunks)}", flush=True)

    client = Ark(base_url=BASE_URL, api_key=API_KEY)

    pages: List[dict] = []
    all_errors: List[str] = []

    prev_sec_idx: Optional[int] = None
    prev_sec_title: Optional[str] = None

    # 为了支持“失败则二分 chunk 重试”，用 while 而不是 for
    i = 0
    part_counter = 0  # 用于 debug 文件命名（即使chunks被插入也不混乱）
    split_count: Dict[int, int] = {}  # original_index -> times

    while i < len(chunks):
        chunk_text = chunks[i]
        part_counter += 1

        # 记录这个 chunk 的“原始序号”（用于限制拆分次数）
        original_key = i
        split_count.setdefault(original_key, 0)

        print(f"[阶段] 3/5 调用模型抽题：第 {i+1}/{len(chunks)} 段（字符数 {len(chunk_text)}）...", flush=True)

        prompt = _build_extract_prompt(
            chapter_name=chapter_name,
            part_no=i + 1,
            part_total=len(chunks),
            prev_section_index=prev_sec_idx,
            prev_section_title=prev_sec_title,
            text=chunk_text,
        )

        resp = client.chat.completions.create(
            model=TEXT_MODEL,
            temperature=0.1,
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
        )


        raw_reply = resp.choices[0].message.content or ""
        _save_debug_text(debug_dir, f"part_{part_counter:03d}_raw.txt", raw_reply)

        # 1) 直接解析
        obj = None
        try:
            obj = _safe_json_loads(raw_reply)
        except Exception as e1:
            # 2) 让模型修复 JSON
            try:
                _save_debug_text(debug_dir, f"part_{part_counter:03d}_parse_error.txt", str(e1))
                obj = _fix_json_via_model(client, raw_reply)
                _save_debug_text(debug_dir, f"part_{part_counter:03d}_fixed.json", json.dumps(obj, ensure_ascii=False, indent=2))
            except Exception as e2:
                _save_debug_text(debug_dir, f"part_{part_counter:03d}_fix_error.txt", str(e2))

                # 3) 仍失败：大概率输出截断或内容过长，自动二分重试
                if len(chunk_text) >= _MIN_SPLIT_CHARS and split_count[original_key] < _MAX_SPLIT_PER_ORIGINAL_CHUNK:
                    split_count[original_key] += 1
                    mid = len(chunk_text) // 2

                    # 尽量在附近找到边界（换行）切一下
                    left = chunk_text[:mid]
                    right = chunk_text[mid:]

                    # 避免切出空段
                    left = left.strip()
                    right = right.strip()

                    if left and right:
                        print("[警告] 本段返回非合法JSON且修复失败，疑似截断/过长，已自动二分重试", flush=True)
                        # 用二分后的两段替换当前段
                        chunks[i:i+1] = [left, right]
                        # 不递增 i，继续处理新的 left
                        continue

                # 无法继续拆分：抛出异常，debug目录已保存原始回复
                raise RuntimeError(
                    f"模型输出JSON解析失败（已写入 debug：{debug_dir.resolve()}）。"
                    f"\n解析错误：{e1}\n修复错误：{e2}"
                )

        # 解析成功：规范化 sections
        secs = obj.get("sections", [])
        if not isinstance(secs, list):
            secs = []

        normalized_secs: List[dict] = []
        for s in secs:
            if not isinstance(s, dict):
                continue
            s = _normalize_section(s, prev_sec_idx, prev_sec_title)
            normalized_secs.append(s)

        # 更新上一段 section（用于下一段沿用）
        if normalized_secs:
            last = normalized_secs[-1]
            prev_sec_idx = last.get("section_index", prev_sec_idx)
            prev_sec_title = last.get("section_title", prev_sec_title)

        page = {
            "_meta": {
                "page_no": len(pages) + 1,  # 输出页号按成功抽题段落计数
                "image": f"{in_path.name}#part{len(pages)+1}",
                "source_type": "docx",
                "source_file": in_path.name,
            },
            "sections": normalized_secs,
            "errors": obj.get("errors", []),
            "confidence": obj.get("confidence", 0.85),
        }
        pages.append(page)

        if isinstance(page["errors"], list):
            all_errors.extend([str(e) for e in page["errors"]])

        # 本段完成，进入下一段
        i += 1

    print("[阶段] 4/5 构建 questions 扁平列表（供 step2 使用）...", flush=True)
    questions_flat: List[Dict[str, Any]] = []
    seq = 0

    for page in pages:
        meta = page.get("_meta", {})
        page_no = meta.get("page_no")
        image = meta.get("image")

        for sec in page.get("sections", []):
            sec_idx = sec.get("section_index")
            sec_title = sec.get("section_title")
            sec_inferred = bool(sec.get("section_inferred", False))

            sec_key = f"IDX:{sec_idx}|{sec_title}"
            for q in sec.get("questions", []):
                seq += 1
                qno = q.get("question_number")
                raw_text = (q.get("raw_text") or "").strip()

                # 兜底含图标记（模型漏标时补一下）
                if not q.get("has_figure", False):
                    if re.search(r"(如图|见图|图\s*\d+)", raw_text):
                        q["has_figure"] = True
                        if not q.get("figure_note"):
                            q["figure_note"] = "题干提到图（仅标记）"

                questions_flat.append({
                    "uid": f"Q{seq:05d}",
                    "chapter": chapter_name,
                    "page_no": page_no,
                    "image": image,
                    "section_index": sec_idx,
                    "section_title": sec_title,
                    "section_key": sec_key,
                    "question_number": qno,
                    "raw_text": raw_text,
                    "has_figure": bool(q.get("has_figure", False)),
                    "figure_note": q.get("figure_note", None),
                    "errors": q.get("errors", []),
                    "confidence": float(q.get("confidence", 0.85)),
                    "section_inferred": sec_inferred,
                })

    chapter_json = {
        "chapter_name": chapter_name,
        "source_type": "docx",
        "source_file": in_path.name,
        "generated_at": __import__("datetime").datetime.now().isoformat(timespec="seconds"),
        "model": TEXT_MODEL,
        "pages": pages,              # step2 用于“原文整理版”
        "questions": questions_flat, # step2 用于“抽样改编”
        "errors": all_errors,
    }

    print("[阶段] 5/5 写出章节总 JSON ...", flush=True)
    out_json_path.write_text(json.dumps(chapter_json, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[完成] 输出：{out_json_path.resolve()}", flush=True)
    print(f"[提示] 若再次失败，请查看 debug：{debug_dir.resolve()}", flush=True)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("用法：python src\\step1_docx_build_chapter_json.py ch01 path\\to\\xxx.docx")
        raise SystemExit(1)
    main(sys.argv[1], sys.argv[2])
