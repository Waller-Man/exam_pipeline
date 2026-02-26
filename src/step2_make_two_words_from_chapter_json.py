import os
import json
import math
import random
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
import time
from dotenv import load_dotenv
from tqdm import tqdm
from volcenginesdkarkruntime import Ark
from docx import Document
from docx.shared import Pt

load_dotenv()

BASE_URL = os.getenv("ARK_BASE_URL", "https://ark.cn-beijing.volces.com/api/v3")
API_KEY = os.getenv("ARK_API_KEY")
TEXT_MODEL = os.getenv("ARK_TEXT_MODEL") or os.getenv("ARK_VISION_MODEL")

# 核心知识点输入过长时，截断到这个字符数（可自行调大/调小）
MAX_KP_INPUT_CHARS = 12000


def safe_json_loads(s: str) -> dict:
    """兜底解析：允许 JSON 前后混入少量文字，截取最外层 {} 再 loads。"""
    try:
        return json.loads(s)
    except json.JSONDecodeError:
        start = s.find("{")
        end = s.rfind("}")
        if start != -1 and end != -1 and end > start:
            return json.loads(s[start:end + 1])
        raise


def section_heading(section_index: Optional[int], section_title: Optional[str]) -> str:
    if section_index is not None:
        return f"大题{section_index}" + (f"：{section_title}" if section_title else "")
    return "大题（未识别编号）" + (f"：{section_title}" if section_title else "")


def has_choice_ABCD(text: str) -> bool:
    """
    粗判：是否是 A/B/C/D 选项题。识别稿里通常有：
    A. / B. / C. / D.
    """
    t = text.replace(" ", "")
    return ("A." in t or "A．" in t) and ("B." in t or "B．" in t) and ("C." in t or "C．" in t) and ("D." in t or "D．" in t)


def build_kp_prompt(chapter_name: str, questions_flat: List[Dict[str, Any]]) -> str:
    """
    核心知识点：复习笔记式概括（不逐题对应），样式仿你给的示例。
    只把题目摘要送进去，并做截断避免过长。
    """
    lines: List[str] = []
    for q in questions_flat:
        # 只取关键信息，避免 prompt 过长
        uid = q.get("uid")
        sec_idx = q.get("section_index")
        sec_title = q.get("section_title")
        qno = q.get("question_number")
        text = (q.get("raw_text") or "").replace("\n", " ")
        text = text[:180]
        lines.append(f"{uid} | {sec_idx}/{sec_title} | 小题{qno}：{text}")

    joined = "\n".join(lines)
    if len(joined) > MAX_KP_INPUT_CHARS:
        joined = joined[:MAX_KP_INPUT_CHARS] + "\n（后续题目摘要已截断）"

    return f"""
你将收到某章节的一批试题摘要。请输出“核心知识点总结（复习笔记风格）”。
不需要逐题对应，只需概括本章主线与高频考点。
必须只输出 JSON（不要输出解释、不要 markdown）。

写作风格（严格仿照示例）：
- 用“一、二、三 …”作为一级标题
- 一级标题下可用“1. 2. 3.”或短横线分条
- 内容像复习笔记：概念/定义 + 核心特征/性质 + 常考点 + 关系主线
- 可以标注“高频考点/必考/核心主线”等
- 不要出现“根据题目/从题干可知”等过程性话术
- 内容可以充实一些，字数在300-400字以内
JSON 结构（必须完全一致）：
{{
  "kp_title": "{chapter_name}核心知识点",
  "kp_text": "一、...\\n\\n二、...\\n..."
}}

章节题目摘要：
{joined}
""".strip()


def build_rewrite_prompt(kp_text: str, q: Dict[str, Any]) -> str:
    """
    改编：题型与原题一致；若原题是A/B/C/D选择题，则新题也保持A/B/C/D。
    输出必须包含：题干+选项（如有）+答案+解析。
    不需要“知识点小标题”。
    """
    uid = q["uid"]
    raw = q["raw_text"]
    is_abcd = has_choice_ABCD(raw)

    type_hint = "保持原题题型与结构"
    if is_abcd:
        type_hint = "原题为A/B/C/D选择题，新题也必须为A/B/C/D四个选项的选择题，且风格/难度相近"

    return f"""
你是一名出题改编助手。给你“本章核心知识点笔记”和一道原题（识别稿）。
请改编出一题同知识域、同风格、难度相近但题面不同的新题。
必须只输出 JSON。

硬性要求：
1) {type_hint}。
2) 不允许与原题高度相似：不要复用原题完整句子；情境/表述/数据应有变化（可以参考结构与数据范围）。
3）题型必须与原题一致，原题为选择题，改编题为选择题，原题为简答题，改变的题目为简答题。
4）如果改编的题目有材料，例如英语的阅读题或材料题，你的改编需要包含这些材料，并输出新题。
5) 输出必须包含：改编后的“题干+选项(如有)”，以及“答案”和“解析”（解析要说明为什么选该答案）。
6) 不要输出答案以外的多余内容，不要输出 markdown,解析不允许输出全英文。

JSON 结构（必须完全一致）：
{{
  "uid": "{uid}",
  "rewritten_text": "string",
  "answer": "string",
  "explanation": "string"
}}

本章核心知识点笔记：
{kp_text}

原题（识别稿）：
{raw}
""".strip()


def ark_chat_json(
    client: Ark,
    prompt: str,
    temperature: float = 0.4,
    max_retries: int = 3,
    stage: str = "model_call",
) -> dict:
    last_err: Exception | None = None

    for attempt in range(1, max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=TEXT_MODEL,
                temperature=temperature,
                messages=[{"role": "user", "content": prompt}],
                response_format={"type": "json_object"},  # ✅ 强制输出 JSON
            )
            content = resp.choices[0].message.content
            return safe_json_loads(content)

        except Exception as e:
            last_err = e
            # 打印当前阶段+重试次数，方便定位
            msg = str(e)
            if attempt < max_retries:
                wait = 1.5 * attempt
                print(f"[警告] {stage} 第{attempt}次返回非JSON/解析失败，{wait:.1f}s后重试：{msg}", flush=True)
                time.sleep(wait)
                continue

            # 最后一次也失败：抛给外层决定“跳题”
            raise RuntimeError(f"{stage} 连续{max_retries}次失败：{msg}") from e

    # 理论上不会走到这里
    raise RuntimeError(f"{stage} failed") from last_err



def choose_rewrite_targets(questions_flat: List[Dict[str, Any]], ratio: float, seed: int) -> List[Dict[str, Any]]:
    # 只允许改编：无图 + 有 uid + 有题干
    eligible = [
        q for q in questions_flat
        if (q.get("uid") and (q.get("raw_text") or "").strip() and (not q.get("has_figure", False)))
    ]

    n = len(eligible)
    if n == 0:
        return []

    k = int(math.floor(n * ratio))
    if ratio > 0 and k == 0:
        k = 1
    k = min(k, n)

    rng = random.Random(seed)
    return rng.sample(eligible, k=k)


def export_doc_raw(chapter_name: str, chapter_json: dict, out_path: Path) -> None:
    """
    文档1：按 pages 输出，最贴近扫描页结构
    """
    doc = Document()
    doc.add_heading(f"{chapter_name}（原文整理版）", level=0)

    pages: List[dict] = chapter_json.get("pages", [])
    # pages 已经按页码输入顺序写入；保险起见按 page_no 排序
    pages_sorted = sorted(pages, key=lambda p: p.get("_meta", {}).get("page_no", 10**9))

    for page in pages_sorted:
        meta = page.get("_meta", {})
        page_no = meta.get("page_no")
        image = meta.get("image")

        doc.add_heading(f"第 {page_no} 页（{image}）", level=1)

        sections = page.get("sections", [])
        for sec in sections:
            sec_idx = sec.get("section_index")
            sec_title = sec.get("section_title")
            doc.add_heading(section_heading(sec_idx, sec_title), level=2)

            for q in sec.get("questions", []):
                qno = q.get("question_number")
                doc.add_heading(f"小题 {qno}", level=3)

                if q.get("has_figure", False):
                    note = q.get("figure_note") or ""
                    doc.add_paragraph(f"[含图] {note}".strip())

                raw_text = q.get("raw_text", "")
                for line in raw_text.splitlines():
                    doc.add_paragraph(line)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def export_doc_kp_and_rewrites(chapter_name: str, kp: dict, rewrites: List[dict], out_path: Path) -> None:
    """
    文档2：核心知识点 + 已改编题（含答案解析）
    """
    doc = Document()
    doc.add_heading(f"{chapter_name}（核心知识点 + 改编题）", level=0)

    doc.add_heading(kp["kp_title"], level=1)
    for line in kp["kp_text"].splitlines():
        doc.add_paragraph(line)

    doc.add_heading("改编题（含答案与解析）", level=1)
    if not rewrites:
        doc.add_paragraph("本章无可改编题（可能因为题目均含图或比例为0）。")
    else:
        for item in rewrites:
            uid = item.get("uid")

            sec_idx = item.get("source_section_index")
            sec_title = item.get("source_section_title")
            qno = item.get("source_question_number")
            src_uid = item.get("source_uid")

            # 标题：给出“原题编号 + uid”
            doc.add_heading(f"改编映射：大题{sec_idx}-小题{qno}（原uid={src_uid}） -> 新uid={uid}", level=2)
            if sec_title:
                doc.add_paragraph(f"所属大题标题：{sec_title}")

            # 原题正文（对照用）
            doc.add_heading("原题（对照）", level=3)
            src_text = item.get("source_raw_text", "") or ""
            for line in src_text.splitlines():
                doc.add_paragraph(line)

            # 改编题正文
            doc.add_heading("改编题", level=3)
            rewritten_text = item.get("rewritten_text", "") or ""
            for line in rewritten_text.splitlines():
                doc.add_paragraph(line)

            # 答案与解析
            doc.add_paragraph(f"答案：{item.get('answer','')}")
            doc.add_paragraph(f"解析：{item.get('explanation','')}")
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main(chapter_name: str, ratio: float = 0.3, seed: int = 2026):
    if not API_KEY or not TEXT_MODEL:
        raise RuntimeError("缺少环境变量：ARK_API_KEY 或 ARK_TEXT_MODEL（请检查 .env）")

    chapter_json_path = Path("out") / chapter_name / f"{chapter_name}.json"
    if not chapter_json_path.exists():
        raise FileNotFoundError(f"找不到：{chapter_json_path.resolve()}（先运行 step1 生成章节总json）")

    print("[阶段] 1/6 读取章节总 JSON...", flush=True)
    chapter_json = json.loads(chapter_json_path.read_text(encoding="utf-8"))

    # 使用 step1 拉平后的 questions（更方便抽样改编）
    questions_flat: List[Dict[str, Any]] = chapter_json.get("questions", [])
    if not questions_flat:
        raise RuntimeError("章节 JSON 中 questions 为空，请检查 step1 输出")

    client = Ark(base_url=BASE_URL, api_key=API_KEY)

    print("[阶段] 2/6 生成核心知识点（概括型，不逐题对应）...", flush=True)
    kp_prompt = build_kp_prompt(chapter_name, questions_flat)
    kp = ark_chat_json(client, kp_prompt, temperature=0.2)

    print(f"[阶段] 3/6 选择改编题：ratio={ratio}（含图题自动跳过）...", flush=True)
    targets = choose_rewrite_targets(questions_flat, ratio=ratio, seed=seed)

    rewrites: List[dict] = []
    if targets:
        print(f"[阶段] 4/6 开始逐题改编：数量={len(targets)}", flush=True)
        for i, q in enumerate(tqdm(targets, desc="改编题"), start=1):
            uid = q.get("uid")
            print(f"[阶段] 4/6 进度 {i}/{len(targets)}：{uid}", flush=True)
            prompt = build_rewrite_prompt(kp["kp_text"], q)
            try:
                out = ark_chat_json(
                    client,
                    prompt,
                    temperature=0.5,
                    max_retries=3,            # 你想重试几次就写几次
                    stage=f"改编题 {uid}",     # 日志里能看见是哪道题
                )
            except Exception as e:
                print(f"[跳过] {uid} 改编失败：{e}", flush=True)
                continue

            rewrites.append({
                **out,
                "source_section_index": q.get("section_index"),
                "source_section_title": q.get("section_title"),
                "source_question_number": q.get("question_number"),
                "source_uid": q.get("uid"),
                "source_raw_text": q.get("raw_text"),
            })

    else:
        print("[阶段] 4/6 无可改编题（可能全是含图题，或ratio=0）", flush=True)

    print("[阶段] 5/6 输出两份 Word...", flush=True)
    out_dir = Path("out") / chapter_name
    doc_raw = out_dir / f"{chapter_name}_原文整理版.docx"
    doc_new = out_dir / f"{chapter_name}_核心知识点与改编题.docx"

    export_doc_raw(chapter_name, chapter_json, doc_raw)
    export_doc_kp_and_rewrites(chapter_name, kp, rewrites, doc_new)

    print("[阶段] 6/6 完成", flush=True)
    print(f"已生成：\n- {doc_raw.resolve()}\n- {doc_new.resolve()}", flush=True)


if __name__ == "__main__":
    # 用法：python src\step2_make_two_words_from_chapter_json.py ch01 0.3 2026
    import sys
    if len(sys.argv) < 2:
        print("用法：python src\\step2_make_two_words_from_chapter_json.py ch01 [ratio] [seed]")
        raise SystemExit(1)

    chapter = sys.argv[1]
    ratio = float(sys.argv[2]) if len(sys.argv) >= 3 else 0.3
    seed = int(sys.argv[3]) if len(sys.argv) >= 4 else 2026
    main(chapter, ratio=ratio, seed=seed)
